import os
import json
import re
import sys
from docx import Document
import fitz  # PyMuPDF
import nltk
from nltk.corpus import words, wordnet
from openai import OpenAI

# Specify the local directory for the NLTK corpus
nltk_data_dir = 'src/nltk_data'
nltk.data.path.append(nltk_data_dir)

# Download 'words' and 'wordnet' corpora if not already present
try:
    nltk.data.find('corpora/words.zip')
except LookupError:
    nltk.download('words', download_dir=nltk_data_dir)

try:
    nltk.data.find('corpora/wordnet.zip')
except LookupError:
    nltk.download('wordnet', download_dir=nltk_data_dir)

api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    raise ValueError("No OpenAI API key found in environment variables.")

client = OpenAI(api_key=api_key)

# Load the set of English words for dictionary check
english_words = set(words.words())

# Load WordNet's word list as an additional general dictionary
wordnet_words = set(wordnet.words())

def is_abbreviation(word):
    # Clean the word for punctuation, but keep internal periods like in "U.S."
    cleaned_word = re.sub(r'[^\w.]', '', word)
    
    # Check against NLTK words corpus
    # if cleaned_word.lower() in english_words:
    #     return False
    
    # # Check against WordNet words corpus
    # if cleaned_word.lower() in wordnet_words:
    #     return False
    
    # Ensure it's not a single letter followed by a period
    if re.fullmatch(r'[A-Z]\.', cleaned_word):
        return False
    
    # Check if it matches the abbreviation pattern and is not too short
    if len(cleaned_word) > 1 and re.fullmatch(r'([A-Z]+\.)+[A-Z]*|[A-Z]{2,}', cleaned_word):
        return True
    
    return False

def get_abbreviations_from_docx(file_path):
    doc = Document(file_path)
    abbreviations = set()
    context = []
    
    for para in doc.paragraphs:
        words = para.text.split()
        context.append(para.text)
        for word in words:
            cleaned_word = re.sub(r'[^\w.]', '', word)
            if is_abbreviation(cleaned_word):
                abbreviations.add(cleaned_word)
    
    return abbreviations, "\n".join(context)

def get_abbreviations_from_pdf(file_path):
    doc = fitz.open(file_path)
    abbreviations = set()
    context = []
    
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text = page.get_text()
        words = text.split()
        context.append(text)
        for word in words:
            cleaned_word = re.sub(r'[^\w.]', '', word)
            if is_abbreviation(cleaned_word):
                abbreviations.add(cleaned_word)
    
    return abbreviations, "\n".join(context)

def refine_abbreviations_with_openai(abbreviations, context):
    refined_abbreviations = {}

    prompt = (
        "Based on the context provided, determine whether the following words are used as "
        "abbreviations or as common words. For each word that is an abbreviation, provide its full form.\n\n"
        f"Context: {context}\n\nWords:\n" + "\n".join(abbreviations)
    )

    response = client.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages=[
            {"role": "user", "content": prompt}
        ],
        temperature=0.0,
        max_tokens=500
    )

    if response.choices:
        response_text = response.choices[0].message.content.strip()
        for line in response_text.split('\n'):
            if ": " in line:
                abbr, meaning = line.split(": ", 1)
                refined_abbreviations[abbr.strip()] = meaning.strip()

    return refined_abbreviations

def extract_context_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_context_from_pdf(file_path):
    doc = fitz.open(file_path)
    full_text = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text = page.get_text()
        full_text.append(text)
    return '\n'.join(full_text)

def main(file_path):
    try:
        if file_path.endswith('.docx'):
            abbreviations, context = get_abbreviations_from_docx(file_path)
        elif file_path.endswith('.pdf'):
            abbreviations, context = get_abbreviations_from_pdf(file_path)
        else:
            raise ValueError("Unsupported file format. Please provide a DOCX or PDF file.")
        
        # Refine abbreviations and get meanings using OpenAI's context understanding
        meanings = refine_abbreviations_with_openai(abbreviations, context)
        
        return meanings
    except Exception as e:
        print(f"Error during processing: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python extractor.py <file_path>', file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    meanings = main(file_path)
    print(json.dumps(meanings))